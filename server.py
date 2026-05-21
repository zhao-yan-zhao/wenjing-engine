import io
import json
import mimetypes
import os
import re
import secrets
import shutil
import subprocess
import threading
import urllib.error
import urllib.request
import uuid
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime
from email.parser import BytesParser
from email.policy import default
from hashlib import pbkdf2_hmac
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import unquote

from docx import Document

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
DATA_DIR = BASE_DIR / "data"
TOOLS_DIR = BASE_DIR / "tools"
USERS_FILE = DATA_DIR / "users.json"
ANALYZER_SCRIPT = TOOLS_DIR / "analyze_aigc_report.py"

MAX_FILE_BYTES = 25 * 1024 * 1024
MAX_AI_CHARS = 12000
AI_CHUNK_SIZE = 10000
AI_CHUNK_OVERLAP = 400
AI_MIN_PARAGRAPH_CHARS = int(os.getenv("AI_MIN_PARAGRAPH_CHARS", "80"))
AI_MAX_TOTAL_CHARS = int(os.getenv("AI_MAX_TOTAL_CHARS", "4200"))
AI_MAX_PARAGRAPHS_PER_JOB = int(os.getenv("AI_MAX_PARAGRAPHS_PER_JOB", "8"))

UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
DATA_DIR.mkdir(exist_ok=True)
TOOLS_DIR.mkdir(exist_ok=True)

JOBS = {}
JOBS_LOCK = threading.Lock()
SESSIONS = {}
SESSIONS_LOCK = threading.Lock()
USERS_LOCK = threading.Lock()


def now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def load_dotenv(path: Path) -> None:
    if not path.exists():
        return

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue

        key, value = line.split("=", 1)
        key = key.strip().lstrip("\ufeff")
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


load_dotenv(BASE_DIR / ".env")

LLM_PROVIDER = os.getenv("LLM_PROVIDER", "openai").strip().lower()
OPENAI_API_URL = os.getenv("OPENAI_API_URL", "https://api.openai.com/v1/responses")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.5")
DEEPSEEK_API_URL = os.getenv("DEEPSEEK_API_URL", "https://api.deepseek.com/chat/completions")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")

ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin123456")


def read_json_file(path: Path, default_value):
    if not path.exists():
        return default_value
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default_value


def write_json_file(path: Path, value) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def hash_password(password: str, salt_hex: str | None = None) -> str:
    if salt_hex is None:
        salt_hex = secrets.token_hex(16)
    salt = bytes.fromhex(salt_hex)
    key = pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 120000)
    return f"{salt_hex}${key.hex()}"


def verify_password(password: str, hashed: str) -> bool:
    if "$" not in hashed:
        return False
    salt_hex, key_hex = hashed.split("$", 1)
    calc = hash_password(password, salt_hex)
    return calc.split("$", 1)[1] == key_hex


def ensure_default_admin() -> None:
    with USERS_LOCK:
        users = read_json_file(USERS_FILE, [])
        if any(u.get("username") == ADMIN_USERNAME for u in users):
            return

        users.append(
            {
                "username": ADMIN_USERNAME,
                "password_hash": hash_password(ADMIN_PASSWORD),
                "role": "admin",
                "created_at": now_str(),
            }
        )
        write_json_file(USERS_FILE, users)


def get_user(username: str):
    with USERS_LOCK:
        users = read_json_file(USERS_FILE, [])
        for user in users:
            if user.get("username") == username:
                return user
    return None


def create_user(username: str, password: str) -> tuple[bool, str]:
    if not re.match(r"^[a-zA-Z0-9_]{3,24}$", username):
        return False, "用户名需为 3-24 位字母、数字或下划线"
    if len(password) < 6:
        return False, "密码至少需要 6 位"

    with USERS_LOCK:
        users = read_json_file(USERS_FILE, [])
        if any(u.get("username") == username for u in users):
            return False, "用户名已存在"

        users.append(
            {
                "username": username,
                "password_hash": hash_password(password),
                "role": "user",
                "created_at": now_str(),
            }
        )
        write_json_file(USERS_FILE, users)
    return True, "注册成功"


def list_users_safe():
    with USERS_LOCK:
        users = read_json_file(USERS_FILE, [])
    return [
        {
            "username": user.get("username", ""),
            "role": user.get("role", "user"),
            "created_at": user.get("created_at", ""),
        }
        for user in users
    ]


def create_session(username: str, role: str) -> str:
    token = secrets.token_urlsafe(32)
    with SESSIONS_LOCK:
        SESSIONS[token] = {
            "username": username,
            "role": role,
            "created_at": now_str(),
        }
    return token


def get_auth_token(headers) -> str | None:
    auth = headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        return auth[7:].strip()
    token = headers.get("X-Auth-Token", "").strip()
    if token:
        return token
    return None


def get_session_user(headers):
    token = get_auth_token(headers)
    if not token:
        return None

    with SESSIONS_LOCK:
        session = SESSIONS.get(token)
        if not session:
            return None
        return {
            "token": token,
            "username": session.get("username", ""),
            "role": session.get("role", "user"),
        }


def destroy_session(token: str) -> None:
    with SESSIONS_LOCK:
        if token in SESSIONS:
            del SESSIONS[token]


def safe_name(filename: str) -> str:
    name = Path(filename).name
    name = re.sub(r"[^\w\-.\u4e00-\u9fff]", "_", name)
    return name[:120] or f"file_{uuid.uuid4().hex[:8]}.txt"


def get_app_version() -> dict:
    commit = os.getenv("RENDER_GIT_COMMIT", "").strip()
    if not commit:
        try:
            commit = subprocess.check_output(
                ["git", "rev-parse", "--short", "HEAD"],
                cwd=BASE_DIR,
                stderr=subprocess.DEVNULL,
                text=True,
            ).strip()
        except Exception:
            commit = "unknown"

    model_name = DEEPSEEK_MODEL if LLM_PROVIDER == "deepseek" else OPENAI_MODEL
    return {
        "commit": commit[:7] if commit and commit != "unknown" else commit,
        "provider": LLM_PROVIDER,
        "model": model_name,
    }


def parse_multipart_form_data(body: bytes, content_type: str) -> tuple[dict, dict]:
    raw_message = (
        f"Content-Type: {content_type}\r\n"
        "MIME-Version: 1.0\r\n\r\n"
    ).encode("utf-8") + body

    message = BytesParser(policy=default).parsebytes(raw_message)
    if not message.is_multipart():
        raise ValueError("请求体不是 multipart/form-data")

    fields = {}
    files = {}
    for part in message.iter_parts():
        disposition = part.get("Content-Disposition", "")
        if "form-data" not in disposition:
            continue

        name = part.get_param("name", header="Content-Disposition")
        if not name:
            continue

        filename = part.get_filename()
        payload = part.get_payload(decode=True) or b""

        if filename is None:
            fields[name] = payload.decode("utf-8", errors="ignore")
        else:
            files[name] = {"filename": filename, "data": payload}

    return fields, files


def parse_json_body(handler) -> dict:
    content_length = int(handler.headers.get("content-length", "0"))
    if content_length <= 0:
        return {}
    raw = handler.rfile.read(content_length)
    if not raw:
        return {}
    return json.loads(raw.decode("utf-8"))


def decode_text_bytes(file_data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return file_data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_data.decode("utf-8", errors="ignore")


def extract_docx_text(file_data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(file_data)) as zf:
        xml_bytes = zf.read("word/document.xml")

    root = ET.fromstring(xml_bytes)
    pieces = []
    for node in root.iter():
        tag = node.tag
        if tag.endswith("}t") and node.text:
            pieces.append(node.text)
        elif tag.endswith("}p"):
            pieces.append("\n")

    text = "".join(pieces)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def file_to_text(file_data: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return decode_text_bytes(file_data).strip()
    if suffix == ".docx":
        return extract_docx_text(file_data)
    if suffix == ".doc":
        raise ValueError("旧版 .doc 暂不支持，请先另存为 .docx")
    return decode_text_bytes(file_data).strip()


def extract_output_text(response_json: dict) -> str:
    top_text = response_json.get("output_text")
    if isinstance(top_text, str) and top_text.strip():
        return top_text.strip()

    chunks = []
    for item in response_json.get("output", []):
        if item.get("type") != "message":
            continue
        for content in item.get("content", []):
            ctype = content.get("type")
            if ctype in {"output_text", "text"} and content.get("text"):
                chunks.append(content["text"])

    merged = "\n".join(chunks).strip()
    if merged:
        return merged
    raise RuntimeError("模型没有返回可用文本")


def sanitize_ai_error(exc: Exception) -> str:
    message = str(exc)
    lowered = message.lower()

    if "incorrect api key" in lowered or "invalid api key" in lowered or "api key" in lowered:
        return "模型服务鉴权失败，请联系管理员更新密钥。"
    if "timeout" in lowered or "timed out" in lowered:
        return "模型服务响应超时，已切换到本地规则降重。"
    if "quota" in lowered or "balance" in lowered:
        return "模型服务额度不足，已切换到本地规则降重。"
    if "401" in lowered or "403" in lowered:
        return "模型服务权限校验失败，已切换到本地规则降重。"

    return "模型服务暂时不可用，已切换到本地规则降重。"


def analyze_text_risk(text: str) -> dict:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    if not normalized:
        return {
            "score": 0,
            "label": "无内容",
            "signals": ["文本为空"],
            "flagged_sentences": [],
        }

    sentences = [s.strip() for s in re.split(r"(?<=[。！？!?；;])", normalized) if s.strip()]
    if not sentences:
        sentences = [normalized]

    total = len(sentences)
    unique = len(set(sentences))
    repeated_ratio = 1 - (unique / total)

    lengths = [len(s) for s in sentences]
    avg_len = sum(lengths) / len(lengths)
    variance = sum((x - avg_len) ** 2 for x in lengths) / len(lengths)
    std_len = variance ** 0.5

    template_phrases = [
        "综上所述",
        "不难发现",
        "值得注意的是",
        "由此可见",
        "因此可以看出",
        "首先",
        "其次",
        "最后",
        "总而言之",
    ]
    template_hits = sum(normalized.count(p) for p in template_phrases)

    connectors = ["因此", "然而", "此外", "同时", "进一步", "总之", "总体来看"]
    connector_hits = sum(normalized.count(c) for c in connectors)

    score = 18
    score += min(32, repeated_ratio * 100 * 0.9)
    score += min(20, max(0, 8 - std_len) * 2.5)
    score += min(18, template_hits * 2)
    score += min(12, connector_hits * 0.8)
    score = int(max(5, min(95, score)))

    if score >= 75:
        label = "高疑似"
    elif score >= 50:
        label = "中疑似"
    else:
        label = "低疑似"

    repeated_sentences = {x for x in sentences if sentences.count(x) > 1}
    sentence_scores = []
    for sentence in sentences:
        sentence_score = 20
        sentence_score += 15 if len(sentence) > 45 else 0
        sentence_score += 20 if any(tp in sentence for tp in template_phrases) else 0
        sentence_score += 15 if sentence in repeated_sentences else 0
        sentence_score += 10 if sentence.count("，") >= 4 else 0
        sentence_scores.append((sentence, min(95, sentence_score)))

    sentence_scores.sort(key=lambda x: x[1], reverse=True)
    flagged = [{"text": s, "score": sc} for s, sc in sentence_scores[:6] if sc >= 45]

    signals = [
        f"句子重复率约 {repeated_ratio * 100:.1f}%",
        f"句长波动标准差 {std_len:.1f}",
        f"模板化短语命中 {template_hits} 次",
    ]

    return {
        "score": score,
        "label": label,
        "signals": signals,
        "flagged_sentences": flagged,
    }


def split_paragraphs(text: str) -> list[str]:
    raw = (text or "").replace("\r\n", "\n")
    parts = re.split(r"\n\s*\n", raw)
    return [p.strip() for p in parts if p.strip()]


def choose_target_paragraphs(paragraphs: list[str], level: str) -> list[int]:
    if not paragraphs:
        return []

    threshold_map = {"标准": 62, "增强": 54, "深度": 46}
    max_targets_map = {"标准": 8, "增强": 16, "深度": 28}

    threshold = threshold_map.get(level, 62)
    max_targets = max_targets_map.get(level, 8)

    scored = []
    for idx, para in enumerate(paragraphs):
        risk = analyze_text_risk(para).get("score", 0)
        scored.append((idx, risk))

    candidates = [x for x in scored if x[1] >= threshold]
    if not candidates:
        top = sorted(scored, key=lambda x: x[1], reverse=True)[:1]
        return [x[0] for x in top]

    candidates.sort(key=lambda x: x[1], reverse=True)
    picked = [idx for idx, _ in candidates[:max_targets]]
    picked.sort()
    return picked


def clamp_target_paragraphs_by_budget(paragraphs: list[str], indexes: list[int]) -> list[int]:
    picked = []
    total_chars = 0

    for idx in indexes:
        para = paragraphs[idx].strip()
        para_len = len(para)
        if para_len < AI_MIN_PARAGRAPH_CHARS:
            continue
        if len(picked) >= AI_MAX_PARAGRAPHS_PER_JOB:
            break
        if picked and total_chars + para_len > AI_MAX_TOTAL_CHARS:
            continue
        if not picked and para_len > AI_MAX_TOTAL_CHARS:
            picked.append(idx)
            break

        picked.append(idx)
        total_chars += para_len

    return picked


def optimize_text_rule(text: str, level: str) -> str:
    replacements = {
        "标准": [("本文", "本研究"), ("我们", "本文"), ("非常", "较为")],
        "增强": [("本文", "本研究"), ("我们", "本文"), ("非常", "较为"), ("所以", "因此"), ("比如", "例如")],
        "深度": [
            ("本文", "本研究"),
            ("我们", "本文"),
            ("非常", "较为"),
            ("所以", "因此"),
            ("比如", "例如"),
            ("很多", "大量"),
            ("这个", "该"),
        ],
    }

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    polished = "\n".join(lines)
    for old, new in replacements.get(level, replacements["标准"]):
        polished = polished.replace(old, new)
    return polished + "\n"


def optimize_text_with_openai(text: str, level: str) -> tuple[str, str]:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("未配置 OPENAI_API_KEY")

    level_rules = {
        "标准": "轻改写，保留结构。",
        "增强": "中度改写，增强衔接。",
        "深度": "深度改写，增强学术表达。",
    }

    payload = {
        "model": OPENAI_MODEL,
        "input": [
            {
                "role": "developer",
                "content": "学术改写助手。只输出改写后的正文。保留原意、数据、引用，不要解释，不要虚构。",
            },
            {
                "role": "user",
                "content": f"强度：{level}\n要求：{level_rules.get(level, level_rules['标准'])}\n{text}",
            },
        ],
    }

    req = urllib.request.Request(
        OPENAI_API_URL,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        method="POST",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
    )

    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            raw = resp.read()
    except urllib.error.HTTPError as http_err:
        err_raw = http_err.read().decode("utf-8", errors="ignore")
        try:
            err_json = json.loads(err_raw)
            message = err_json.get("error", {}).get("message") or err_raw
        except json.JSONDecodeError:
            message = err_raw or str(http_err)
        raise RuntimeError(f"OpenAI 接口错误：{message}") from http_err
    except Exception as exc:
        raise RuntimeError(f"OpenAI 请求失败：{exc}") from exc

    response_json = json.loads(raw.decode("utf-8"))
    return extract_output_text(response_json), OPENAI_MODEL


def optimize_text_with_deepseek(text: str, level: str) -> tuple[str, str]:
    api_key = os.getenv("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("未配置 DEEPSEEK_API_KEY")

    level_rules = {
        "标准": "轻改写，保留结构。",
        "增强": "中度改写，增强衔接。",
        "深度": "深度改写，增强学术表达。",
    }

    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {
                "role": "system",
                "content": "学术改写助手。只输出改写后的正文。保留原意、数据、引用，不要解释，不要虚构。",
            },
            {
                "role": "user",
                "content": f"强度：{level}\n要求：{level_rules.get(level, level_rules['标准'])}\n{text}",
            },
        ],
        "stream": False,
    }

    req = urllib.request.Request(
        DEEPSEEK_API_URL,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        method="POST",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
    )

    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            raw = resp.read()
    except urllib.error.HTTPError as http_err:
        err_raw = http_err.read().decode("utf-8", errors="ignore")
        try:
            err_json = json.loads(err_raw)
            message = err_json.get("error", {}).get("message") or err_raw
        except json.JSONDecodeError:
            message = err_raw or str(http_err)
        raise RuntimeError(f"DeepSeek 接口错误：{message}") from http_err
    except Exception as exc:
        raise RuntimeError(f"DeepSeek 请求失败：{exc}") from exc

    response_json = json.loads(raw.decode("utf-8"))
    choices = response_json.get("choices", [])
    if not choices:
        raise RuntimeError("DeepSeek 没有返回内容")
    content = choices[0].get("message", {}).get("content", "")
    if not isinstance(content, str) or not content.strip():
        raise RuntimeError("DeepSeek 没有返回内容")
    return content.strip(), DEEPSEEK_MODEL


def optimize_text_with_provider(text: str, level: str) -> tuple[str, str]:
    if not text.strip():
        raise RuntimeError("文本为空")

    if LLM_PROVIDER == "deepseek":
        out, model = optimize_text_with_deepseek(text, level)
        return out, f"deepseek:{model}"
    if LLM_PROVIDER == "openai":
        out, model = optimize_text_with_openai(text, level)
        return out, f"openai:{model}"
    raise RuntimeError(f"不支持的 LLM_PROVIDER: {LLM_PROVIDER}")


def split_text_into_chunks(text: str, size: int = AI_CHUNK_SIZE, overlap: int = AI_CHUNK_OVERLAP) -> list[str]:
    normalized = text.replace("\r\n", "\n").strip()
    if len(normalized) <= size:
        return [normalized]

    chunks = []
    start = 0
    length = len(normalized)

    while start < length:
        end = min(start + size, length)
        if end < length:
            cut = normalized.rfind("\n", start + int(size * 0.6), end)
            if cut == -1:
                cut = normalized.rfind("。", start + int(size * 0.6), end)
            if cut != -1:
                end = cut + 1

        part = normalized[start:end].strip()
        if part:
            chunks.append(part)

        if end >= length:
            break
        start = max(0, end - overlap)

    return chunks


def optimize_text_with_provider_chunked(text: str, level: str) -> tuple[str, str]:
    chunks = split_text_into_chunks(text)
    result_parts = []
    engine_name = ""

    for idx, chunk in enumerate(chunks, start=1):
        if len(chunk) > MAX_AI_CHARS:
            raise RuntimeError(f"分段后仍超限：第 {idx} 段 {len(chunk)} 字符")
        optimized, engine_name = optimize_text_with_provider(chunk, level)
        result_parts.append(optimized.strip())

    return "\n\n".join(result_parts).strip(), engine_name


def build_visual_report(before: dict, after: dict, rewrite_targets: int, estimated_ai_chars: int, extra: dict | None = None) -> dict:
    report = {
        "before": before,
        "after": after,
        "drop": max(0, before.get("score", 0) - after.get("score", 0)),
        "rewrite_targets": rewrite_targets,
        "estimated_ai_chars": estimated_ai_chars,
    }
    if extra:
        report.update(extra)
    return report


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def parse_target_paragraphs_tsv(path: Path) -> list[int]:
    indexes = []
    if not path.exists():
        return indexes

    lines = path.read_text(encoding="utf-8").splitlines()[1:]
    for line in lines:
        parts = line.split("\t", 4)
        if parts and parts[0].isdigit():
            indexes.append(int(parts[0]))
    return indexes


def run_report_analyzer(docx_path: Path, report_path: Path, work_dir: Path) -> dict:
    if not ANALYZER_SCRIPT.exists():
        raise RuntimeError(f"缺少分析脚本：{ANALYZER_SCRIPT}")

    cmd = [
        "python",
        str(ANALYZER_SCRIPT),
        "--docx",
        str(docx_path),
        "--report",
        str(report_path),
        "--out",
        str(work_dir),
        "--levels",
        "high,medium",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="ignore")
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or "报告分析失败")
    return read_json_file(work_dir / "stats.json", {})


def rewrite_docx_with_report(job_id: str, source_docx_path: Path, report_path: Path, level: str, saved_output: Path) -> dict:
    work_dir = BASE_DIR / "work" / f"job_{job_id}"
    work_dir.mkdir(parents=True, exist_ok=True)

    stats = run_report_analyzer(source_docx_path, report_path, work_dir)
    target_indexes = parse_target_paragraphs_tsv(work_dir / "target_paragraphs.tsv")

    doc = Document(str(source_docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    chosen_indexes = clamp_target_paragraphs_by_budget(paragraphs, target_indexes)

    before_text = "\n\n".join(p for p in paragraphs if p.strip())
    before_risk = analyze_text_risk(before_text)
    estimated_ai_chars = sum(len(paragraphs[idx].strip()) for idx in chosen_indexes if idx < len(paragraphs))
    update_job(
        job_id,
        progress=18,
        message=f"报告已解析，命中 {len(target_indexes)} 个段落，准备定向改写...",
        rewrite_targets=len(chosen_indexes),
        report_summary=stats,
        aigc_before=before_risk,
        estimated_ai_chars=estimated_ai_chars,
    )

    cache = {}
    for i, idx in enumerate(chosen_indexes, start=1):
        para_text = doc.paragraphs[idx].text.strip()
        if len(para_text) < AI_MIN_PARAGRAPH_CHARS:
            continue

        cache_key = re.sub(r"\s+", " ", para_text)
        if cache_key in cache:
            rewritten = cache[cache_key]
        elif len(para_text) <= MAX_AI_CHARS:
            rewritten, _ = optimize_text_with_provider(para_text, level)
            cache[cache_key] = rewritten
        else:
            rewritten, _ = optimize_text_with_provider_chunked(para_text, level)
            cache[cache_key] = rewritten

        replace_paragraph_text(doc.paragraphs[idx], rewritten.strip())
        ratio = i / max(len(chosen_indexes), 1)
        progress = 22 + int(ratio * 63)
        update_job(job_id, progress=min(progress, 90), message=f"报告定向改写中：第 {i}/{len(chosen_indexes)} 段")

    doc.save(str(saved_output))

    out_doc = Document(str(saved_output))
    after_text = "\n\n".join(p.text.strip() for p in out_doc.paragraphs if p.text.strip())
    after_risk = analyze_text_risk(after_text)
    drop_value = max(0, before_risk.get("score", 0) - after_risk.get("score", 0))

    return {
        "engine": f"{LLM_PROVIDER}:report-guided",
        "notice": f"按报告定向改写 {len(chosen_indexes)} 个段落",
        "bytes": saved_output.stat().st_size,
        "aigc_before": before_risk,
        "aigc_after": after_risk,
        "aigc_drop": drop_value,
        "rewrite_targets": len(chosen_indexes),
        "estimated_ai_chars": estimated_ai_chars,
        "visual_report": build_visual_report(
            before_risk,
            after_risk,
            len(chosen_indexes),
            estimated_ai_chars,
            {
                "report_counts": stats.get("report_counts", {}),
                "mapped_items": stats.get("mapped_items", 0),
                "unique_target_paragraphs": stats.get("unique_target_paragraphs", 0),
                "report_type": stats.get("report_type", ""),
            },
        ),
    }


def process_plain_text(job_id: str, source_text: str, level: str, saved_output: Path) -> dict:
    before_risk = analyze_text_risk(source_text)
    update_job(job_id, status="processing", progress=10, message="正在分析 AIGC 风险...", aigc_before=before_risk)

    paragraphs = split_paragraphs(source_text)
    target_indexes = choose_target_paragraphs(paragraphs, level)
    target_indexes = clamp_target_paragraphs_by_budget(paragraphs, target_indexes)
    estimated_ai_chars = sum(len(paragraphs[idx].strip()) for idx in target_indexes) if paragraphs else 0
    update_job(
        job_id,
        progress=16,
        message=f"已锁定 {len(target_indexes)} 个高风险段落，准备降重...",
        rewrite_targets=len(target_indexes),
        estimated_ai_chars=estimated_ai_chars,
    )

    engine = "rule-fallback"
    notice = ""
    rewritten = list(paragraphs)
    cache = {}
    skipped_short = 0

    try:
        if not target_indexes:
            target_indexes = [0] if paragraphs else []

        for i, para_idx in enumerate(target_indexes, start=1):
            para_text = rewritten[para_idx]
            para_key = re.sub(r"\s+", " ", para_text).strip()

            if len(para_text.strip()) < AI_MIN_PARAGRAPH_CHARS:
                skipped_short += 1
                continue

            if para_key in cache:
                new_text, engine = cache[para_key]
            elif len(para_text) <= MAX_AI_CHARS:
                new_text, engine = optimize_text_with_provider(para_text, level)
                cache[para_key] = (new_text, engine)
            else:
                new_text, engine = optimize_text_with_provider_chunked(para_text, level)
                cache[para_key] = (new_text, engine)

            rewritten[para_idx] = new_text.strip()
            ratio = i / max(len(target_indexes), 1)
            pct = 20 + int(ratio * 65)
            update_job(job_id, progress=min(pct, 90), message=f"降重处理中：第 {i}/{len(target_indexes)} 段")

        optimized_body = "\n\n".join(rewritten).strip()
        if len(target_indexes) > 1:
            notice = f"采用定向降重：共改写 {len(target_indexes)} 个风险段落"
        if skipped_short:
            suffix = f"；跳过 {skipped_short} 个过短段落以节省 token"
            notice = f"{notice}{suffix}" if notice else suffix.lstrip("；")
    except Exception as ai_exc:
        if rewritten:
            for idx in target_indexes:
                rewritten[idx] = optimize_text_rule(rewritten[idx], level).strip()
            optimized_body = "\n\n".join(rewritten).strip()
        else:
            optimized_body = optimize_text_rule(source_text, level).strip()

        engine = "rule-fallback"
        notice = sanitize_ai_error(ai_exc)
        update_job(job_id, progress=86, message="模型服务不可用，正在回退规则引擎...")

    after_risk = analyze_text_risk(optimized_body)
    drop_value = max(0, before_risk.get("score", 0) - after_risk.get("score", 0))

    result_text = (
        "【文净引擎处理结果】\n"
        f"处理时间：{now_str()}\n"
        f"优化强度：{level}\n"
        f"处理引擎：{engine}\n"
        f"AIGC疑似率(估计)：{before_risk.get('score', 0)} -> {after_risk.get('score', 0)}（下降 {drop_value}）\n"
        "----------------------------------------\n"
        f"{optimized_body}\n"
    )
    saved_output.write_text(result_text, encoding="utf-8")

    return {
        "engine": engine,
        "notice": notice,
        "bytes": saved_output.stat().st_size,
        "aigc_before": before_risk,
        "aigc_after": after_risk,
        "aigc_drop": drop_value,
        "rewrite_targets": len(target_indexes),
        "estimated_ai_chars": estimated_ai_chars,
        "visual_report": build_visual_report(before_risk, after_risk, len(target_indexes), estimated_ai_chars),
    }


def create_job(username: str, level: str, filename: str) -> str:
    job_id = uuid.uuid4().hex[:12]
    with JOBS_LOCK:
        JOBS[job_id] = {
            "job_id": job_id,
            "status": "queued",
            "progress": 0,
            "message": "任务已创建",
            "engine": "",
            "notice": "",
            "download_url": "",
            "output_name": "",
            "bytes": 0,
            "error": "",
            "aigc_before": None,
            "aigc_after": None,
            "aigc_drop": None,
            "rewrite_targets": 0,
            "estimated_ai_chars": 0,
            "visual_report": None,
            "report_summary": None,
            "submitted_by": username,
            "level": level,
            "source_filename": filename,
            "created_at": now_str(),
            "updated_at": now_str(),
        }
    return job_id


def update_job(job_id: str, **kwargs) -> None:
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if not job:
            return
        job.update(kwargs)
        job["updated_at"] = now_str()


def get_job(job_id: str):
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        return dict(job) if job else None


def list_jobs_for_user(username: str):
    with JOBS_LOCK:
        jobs = [dict(v) for v in JOBS.values() if v.get("submitted_by") == username]
    jobs.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return jobs


def list_all_jobs():
    with JOBS_LOCK:
        jobs = [dict(v) for v in JOBS.values()]
    jobs.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return jobs


def process_document_job(job_id: str, source_text: str, level: str, saved_output: Path, source_docx_path: Path | None = None, report_path: Path | None = None) -> None:
    try:
        if source_docx_path is not None and report_path is not None:
            result = rewrite_docx_with_report(job_id, source_docx_path, report_path, level, saved_output)
        else:
            result = process_plain_text(job_id, source_text, level, saved_output)

        update_job(
            job_id,
            status="completed",
            progress=100,
            message="处理完成",
            engine=result["engine"],
            notice=result["notice"],
            download_url=f"/api/download/{saved_output.name}",
            output_name=saved_output.name,
            bytes=result["bytes"],
            aigc_before=result["aigc_before"],
            aigc_after=result["aigc_after"],
            aigc_drop=result["aigc_drop"],
            rewrite_targets=result["rewrite_targets"],
            estimated_ai_chars=result["estimated_ai_chars"],
            visual_report=result["visual_report"],
        )
    except Exception as exc:
        update_job(job_id, status="failed", progress=100, message="处理失败", error=str(exc))


class AppHandler(SimpleHTTPRequestHandler):
    server_version = "WenjingEngine/3.1"

    def do_POST(self):
        if self.path == "/api/register":
            self.handle_register()
            return
        if self.path == "/api/login":
            self.handle_login()
            return
        if self.path == "/api/logout":
            self.handle_logout()
            return
        if self.path == "/api/process":
            self.handle_process()
            return
        if self.path == "/api/aigc/check":
            self.handle_aigc_check()
            return
        self.send_error(HTTPStatus.NOT_FOUND, "Not Found")

    def do_GET(self):
        if self.path == "/healthz":
            self.respond_json(HTTPStatus.OK, {"ok": True})
            return
        if self.path == "/api/version":
            self.respond_json(HTTPStatus.OK, get_app_version())
            return
        if self.path == "/api/me":
            self.handle_me()
            return
        if self.path == "/api/my/jobs":
            self.handle_my_jobs()
            return
        if self.path == "/api/admin/users":
            self.handle_admin_users()
            return
        if self.path == "/api/admin/jobs":
            self.handle_admin_jobs()
            return
        if self.path == "/api/admin/stats":
            self.handle_admin_stats()
            return
        if self.path.startswith("/api/job/"):
            self.handle_job_status()
            return
        if self.path.startswith("/api/download/"):
            self.handle_download()
            return
        super().do_GET()

    def current_user(self):
        return get_session_user(self.headers)

    def require_user(self):
        user = self.current_user()
        if not user:
            self.respond_json(HTTPStatus.UNAUTHORIZED, {"error": "请先登录"})
            return None
        return user

    def require_admin(self):
        user = self.require_user()
        if not user:
            return None
        if user.get("role") != "admin":
            self.respond_json(HTTPStatus.FORBIDDEN, {"error": "需要管理员权限"})
            return None
        return user

    def handle_register(self):
        try:
            payload = parse_json_body(self)
            username = str(payload.get("username", "")).strip()
            password = str(payload.get("password", ""))
            ok, msg = create_user(username, password)
            if not ok:
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": msg})
                return
            self.respond_json(HTTPStatus.OK, {"message": msg})
        except Exception as exc:
            self.respond_json(HTTPStatus.INTERNAL_SERVER_ERROR, {"error": f"注册失败：{exc}"})

    def handle_login(self):
        try:
            payload = parse_json_body(self)
            username = str(payload.get("username", "")).strip()
            password = str(payload.get("password", ""))

            user = get_user(username)
            if not user or not verify_password(password, user.get("password_hash", "")):
                self.respond_json(HTTPStatus.UNAUTHORIZED, {"error": "用户名或密码错误"})
                return

            token = create_session(username, user.get("role", "user"))
            self.respond_json(
                HTTPStatus.OK,
                {
                    "message": "登录成功",
                    "token": token,
                    "user": {"username": username, "role": user.get("role", "user")},
                },
            )
        except Exception as exc:
            self.respond_json(HTTPStatus.INTERNAL_SERVER_ERROR, {"error": f"登录失败：{exc}"})

    def handle_logout(self):
        user = self.require_user()
        if not user:
            return
        destroy_session(user.get("token", ""))
        self.respond_json(HTTPStatus.OK, {"message": "已退出登录"})

    def handle_me(self):
        user = self.require_user()
        if not user:
            return
        self.respond_json(HTTPStatus.OK, {"username": user.get("username"), "role": user.get("role")})

    def handle_aigc_check(self):
        user = self.require_user()
        if not user:
            return

        try:
            payload = parse_json_body(self)
            text = str(payload.get("text", "")).strip()
            if not text:
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "text 不能为空"})
                return
            self.respond_json(HTTPStatus.OK, analyze_text_risk(text))
        except Exception as exc:
            self.respond_json(HTTPStatus.INTERNAL_SERVER_ERROR, {"error": f"检测失败：{exc}"})

    def handle_process(self):
        user = self.require_user()
        if not user:
            return

        try:
            content_type = self.headers.get("content-type", "")
            if "multipart/form-data" not in content_type.lower():
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "请求格式必须是 multipart/form-data"})
                return

            content_length = int(self.headers.get("content-length", "0"))
            if content_length > MAX_FILE_BYTES + (1024 * 512):
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "文件过大，最大 25MB"})
                return

            body = self.rfile.read(content_length)
            fields, files = parse_multipart_form_data(body, content_type)

            file_item = files.get("file")
            if file_item is None or not file_item.get("filename"):
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "请上传待处理文档"})
                return

            level = fields.get("level", "标准")
            filename = safe_name(file_item["filename"])
            file_data = file_item["data"]
            report_item = files.get("report")

            if not file_data:
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "文件内容为空"})
                return
            if len(file_data) > MAX_FILE_BYTES:
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "文件过大，最大 25MB"})
                return

            upload_id = uuid.uuid4().hex[:12]
            saved_input = UPLOAD_DIR / f"{upload_id}_{filename}"
            saved_output = OUTPUT_DIR / f"{upload_id}_optimized.txt"
            report_path = None

            saved_input.write_bytes(file_data)

            if report_item and report_item.get("filename"):
                report_name = safe_name(report_item["filename"])
                report_path = UPLOAD_DIR / f"{upload_id}_{report_name}"
                report_path.write_bytes(report_item["data"])

            try:
                source_text = file_to_text(file_data, filename)
            except Exception as parse_exc:
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": f"文档解析失败：{parse_exc}"})
                return

            if not source_text:
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "未提取到可处理文本内容"})
                return

            source_docx_path = saved_input if saved_input.suffix.lower() == ".docx" else None
            if report_path is not None and source_docx_path is None:
                self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "带报告的精准改写目前仅支持 DOCX 文档"})
                return
            if report_path is not None:
                saved_output = OUTPUT_DIR / f"{upload_id}_optimized.docx"

            job_id = create_job(user.get("username", ""), level, filename)
            update_job(job_id, progress=6, message="任务已入队，等待处理...")

            worker = threading.Thread(
                target=process_document_job,
                args=(job_id, source_text, level, saved_output, source_docx_path, report_path),
                daemon=True,
            )
            worker.start()

            self.respond_json(HTTPStatus.ACCEPTED, {"message": "任务已提交", "job_id": job_id, "status": "queued", "progress": 6})
        except Exception as exc:
            self.respond_json(HTTPStatus.INTERNAL_SERVER_ERROR, {"error": f"服务端异常：{exc}"})

    def handle_job_status(self):
        user = self.require_user()
        if not user:
            return

        job_id = unquote(self.path.replace("/api/job/", "", 1)).strip()
        if not job_id:
            self.respond_json(HTTPStatus.BAD_REQUEST, {"error": "缺少 job_id"})
            return

        job = get_job(job_id)
        if not job:
            self.respond_json(HTTPStatus.NOT_FOUND, {"error": "任务不存在"})
            return

        if user.get("role") != "admin" and job.get("submitted_by") != user.get("username"):
            self.respond_json(HTTPStatus.FORBIDDEN, {"error": "无权访问该任务"})
            return

        self.respond_json(HTTPStatus.OK, job)

    def handle_my_jobs(self):
        user = self.require_user()
        if not user:
            return
        self.respond_json(HTTPStatus.OK, {"items": list_jobs_for_user(user.get("username", ""))[:20]})

    def handle_admin_users(self):
        if not self.require_admin():
            return
        self.respond_json(HTTPStatus.OK, {"items": list_users_safe()})

    def handle_admin_jobs(self):
        if not self.require_admin():
            return
        self.respond_json(HTTPStatus.OK, {"items": list_all_jobs()[:100]})

    def handle_admin_stats(self):
        if not self.require_admin():
            return

        users = list_users_safe()
        jobs = list_all_jobs()
        completed = [j for j in jobs if j.get("status") == "completed"]
        avg_drop = 0
        if completed:
            avg_drop = round(sum(j.get("aigc_drop") or 0 for j in completed) / len(completed), 2)

        self.respond_json(
            HTTPStatus.OK,
            {
                "users": len(users),
                "jobs_total": len(jobs),
                "jobs_completed": len(completed),
                "avg_aigc_drop": avg_drop,
            },
        )

    def handle_download(self):
        user = self.require_user()
        if not user:
            return

        token = unquote(self.path.replace("/api/download/", "", 1)).strip()
        if not token:
            self.send_error(HTTPStatus.BAD_REQUEST, "Bad Request")
            return

        file_path = OUTPUT_DIR / Path(token).name
        if not file_path.exists() or not file_path.is_file():
            self.send_error(HTTPStatus.NOT_FOUND, "File Not Found")
            return

        related_job = None
        with JOBS_LOCK:
            for value in JOBS.values():
                if value.get("output_name") == file_path.name:
                    related_job = dict(value)
                    break

        if related_job and user.get("role") != "admin":
            if related_job.get("submitted_by") != user.get("username"):
                self.respond_json(HTTPStatus.FORBIDDEN, {"error": "无权下载该文件"})
                return

        mime_type, _ = mimetypes.guess_type(file_path.name)
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", mime_type or "application/octet-stream")
        self.send_header("Content-Disposition", f'attachment; filename="{file_path.name}"')
        self.send_header("Content-Length", str(file_path.stat().st_size))
        self.end_headers()

        with file_path.open("rb") as f:
            shutil.copyfileobj(f, self.wfile)

    def respond_json(self, status: int, payload: dict):
        raw = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def log_message(self, format, *args):
        return


def run(host: str | None = None, port: int | None = None):
    ensure_default_admin()
    host = host or os.getenv("HOST", "0.0.0.0")
    port = int(port or os.getenv("PORT", "8000"))
    os.chdir(BASE_DIR)
    server = ThreadingHTTPServer((host, port), AppHandler)
    print(f"Server running at http://{host}:{port}")
    print(f"Provider: {LLM_PROVIDER}")
    print("Press Ctrl+C to stop")
    server.serve_forever()


if __name__ == "__main__":
    run()
