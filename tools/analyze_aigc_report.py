#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Analyze AIGC reports and map flagged snippets to DOCX paragraphs."""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
from collections import Counter, defaultdict
from difflib import SequenceMatcher
from pathlib import Path
from zipfile import ZipFile

from docx import Document

REPORT_EXT_ORDER = [".html", ".htm", ".docx", ".pdf"]


def resolve_report_path(path: Path) -> Path:
    if path.is_file():
        return path
    if not path.is_dir():
        raise FileNotFoundError(path)
    files = [p for p in path.rglob("*") if p.is_file()]
    for ext in REPORT_EXT_ORDER:
        candidates = [p for p in files if p.suffix.lower() == ext]
        if candidates:
            candidates.sort(key=lambda p: (len(p.parts), len(p.name), str(p).lower()))
            return candidates[0]
    raise FileNotFoundError(f"No supported report file found in {path}")


def read_text(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "gb18030", "gbk"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            pass
    return raw.decode("utf-8", errors="replace")


def clean(s: str) -> str:
    return re.sub(r"\s+", "", s or "")


def classify_rgb(r: int, g: int, b: int) -> str | None:
    if r >= 150 and g <= 120 and b <= 120 and r >= g + 50 and r >= b + 50:
        return "high"
    if r >= 170 and 70 <= g <= 230 and b <= 140 and r >= b + 50:
        return "medium"
    if r >= 90 and b >= 110 and g <= 140:
        return "low"
    return None


def classify_hex_color(value: str | None) -> str | None:
    if not value:
        return None
    value = value.strip().replace("#", "")
    if len(value) != 6:
        return None
    try:
        r = int(value[0:2], 16)
        g = int(value[2:4], 16)
        b = int(value[4:6], 16)
    except ValueError:
        return None
    return classify_rgb(r, g, b)


def classify_highlight(highlight) -> str | None:
    if highlight is None:
        return None
    name = str(highlight).upper()
    if "RED" in name:
        return "high"
    if "YELLOW" in name or "GOLD" in name:
        return "medium"
    if "VIOLET" in name or "PURPLE" in name or "PINK" in name:
        return "low"
    return None


def merge_items(chunks: list[tuple[str, str]]) -> list[tuple[str, str]]:
    merged: list[tuple[str, str]] = []
    cur_level = None
    cur_text: list[str] = []
    for level, text in chunks:
        text = clean(text)
        if not text:
            continue
        if level == cur_level:
            cur_text.append(text)
        else:
            if cur_level and cur_text:
                merged.append((cur_level, "".join(cur_text)))
            cur_level = level
            cur_text = [text]
    if cur_level and cur_text:
        merged.append((cur_level, "".join(cur_text)))
    return merged


def extract_html_items(report: Path) -> tuple[list[tuple[str, str]], list[str]]:
    text = read_text(report)
    items: list[tuple[str, str]] = []
    pattern = r"<em\s+class=['\"](high|medium|low)['\"][^>]*>(.*?)</em>"
    for match in re.finditer(pattern, text, re.S | re.I):
        level = match.group(1).lower()
        body = re.sub(r"<[^>]+>", "", match.group(2))
        body = html.unescape(body).strip()
        body = clean(body)
        if body:
            items.append((level, body))
    return items, []


def extract_docx_report_items(report: Path) -> tuple[list[tuple[str, str]], list[str]]:
    doc = Document(report)
    chunks: list[tuple[str, str]] = []
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text or ""
            level = None
            color = run.font.color
            if color is not None and color.rgb is not None:
                level = classify_hex_color(str(color.rgb))
            if level is None:
                level = classify_highlight(run.font.highlight_color)
            if level:
                chunks.append((level, text))
        chunks.append(("__break__", ""))
    chunks = [(level, text) for level, text in chunks if level != "__break__"]
    return merge_items(chunks), []


def extract_pdf_report_items(report: Path, out_dir: Path) -> tuple[list[tuple[str, str]], list[str]]:
    warnings: list[str] = []
    try:
        import fitz  # type: ignore
    except Exception:
        warnings.append("PDF 彩色检测需要 PyMuPDF，当前仅支持 HTML/DOCX 报告的自动定位。")
        return [], warnings

    doc = fitz.open(str(report))
    chunks: list[tuple[str, str]] = []
    for page in doc:
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    color = int(span.get("color", 0))
                    r = (color >> 16) & 255
                    g = (color >> 8) & 255
                    b = color & 255
                    level = classify_rgb(r, g, b)
                    if level:
                        chunks.append((level, text))
                chunks.append(("__break__", ""))
    chunks = [(level, text) for level, text in chunks if level != "__break__"]
    return merge_items(chunks), warnings


def extract_report_items(report: Path, out_dir: Path) -> tuple[list[tuple[str, str]], list[str], str]:
    suffix = report.suffix.lower()
    if suffix in (".html", ".htm"):
        items, warnings = extract_html_items(report)
        return items, warnings, "html"
    if suffix == ".docx":
        items, warnings = extract_docx_report_items(report)
        return items, warnings, "docx"
    if suffix == ".pdf":
        items, warnings = extract_pdf_report_items(report, out_dir)
        return items, warnings, "pdf"
    raise ValueError(f"Unsupported report type: {report}")


def doc_stats(doc: Document) -> dict[str, int]:
    return {
        "paragraphs": len(doc.paragraphs),
        "non_empty": sum(1 for p in doc.paragraphs if p.text.strip()),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "para_chars": sum(len(p.text.strip()) for p in doc.paragraphs if p.text.strip()),
        "table_chars": sum(
            len(cell.text.strip())
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        ),
    }


def media_hashes(path: Path) -> dict[str, str]:
    with ZipFile(path) as zf:
        return {
            name: hashlib.md5(zf.read(name)).hexdigest()
            for name in zf.namelist()
            if name.startswith("word/media/")
        }


def map_items_to_paragraphs(items: list[tuple[int, str, str]], doc: Document, min_ratio: float):
    paras = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    mapped = []
    for item_no, level, snippet in items:
        cs = clean(snippet)
        best_idx, best_score, best_para = -1, 0.0, ""
        for idx, para in paras:
            cp = clean(para)
            if cs and (cs in cp or cp in cs):
                best_idx = idx
                best_score = 1.0 + min(len(cs), len(cp)) / max(len(cs), len(cp))
                best_para = para
                break
        if best_idx < 0:
            for idx, para in paras:
                score = SequenceMatcher(None, cs, clean(para)).ratio()
                if score > best_score:
                    best_idx, best_score, best_para = idx, score, para
        if best_score >= min_ratio:
            mapped.append((item_no, level, snippet, best_idx, best_score, best_para))
    return mapped


def write_context(path: Path, doc: Document, targets: dict[int, list[tuple]], window: int) -> None:
    lines: list[str] = []
    for idx in sorted(targets):
        lines.append("=" * 30 + f" TARGET {idx} " + "=" * 30)
        hit_summary = "; ".join(f"{h[1]}#{h[0]} len={len(h[2])} score={h[4]:.3f}" for h in targets[idx])
        lines.append(f"hits: {hit_summary}")
        for j in range(max(0, idx - window), min(len(doc.paragraphs), idx + window + 1)):
            text = doc.paragraphs[j].text.strip()
            if not text:
                continue
            prefix = ">>" if j == idx else "  "
            lines.append(f"{prefix} IDX {j} len={len(text)}")
            lines.append(text)
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--report", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--levels", default="high,medium")
    parser.add_argument("--context-window", type=int, default=2)
    parser.add_argument("--min-ratio", type=float, default=0.45)
    args = parser.parse_args()

    args.out.mkdir(parents=True, exist_ok=True)
    report_path = resolve_report_path(args.report)
    levels = {x.strip().lower() for x in args.levels.split(",") if x.strip()}
    all_items, warnings, report_type = extract_report_items(report_path, args.out)
    numbered = [(i, level, text) for i, (level, text) in enumerate(all_items, 1)]
    targets = [(i, level, text) for i, level, text in numbered if level in levels]

    doc = Document(args.docx)
    mapped = map_items_to_paragraphs(targets, doc, args.min_ratio)
    by_idx: dict[int, list[tuple]] = defaultdict(list)
    for row in mapped:
        by_idx[row[3]].append(row)

    with (args.out / "flagged_items.tsv").open("w", encoding="utf-8") as f:
        f.write("item_no\tlevel\tlength\tsnippet\n")
        for i, level, text in numbered:
            f.write(f"{i}\t{level}\t{len(text)}\t{text}\n")

    with (args.out / "target_paragraphs.tsv").open("w", encoding="utf-8") as f:
        f.write("paragraph_idx\thit_count\tlevels\tpara_length\tparagraph\n")
        for idx in sorted(by_idx):
            para = doc.paragraphs[idx].text.strip()
            level_summary = ",".join(sorted(Counter(row[1] for row in by_idx[idx]).keys()))
            f.write(f"{idx}\t{len(by_idx[idx])}\t{level_summary}\t{len(para)}\t{para}\n")

    write_context(args.out / "context.txt", doc, by_idx, args.context_window)

    stats = {
        "resolved_report": str(report_path),
        "report_type": report_type,
        "warnings": warnings,
        "report_counts": dict(Counter(level for level, _ in all_items)),
        "target_levels": sorted(levels),
        "target_items": len(targets),
        "mapped_items": len(mapped),
        "unique_target_paragraphs": len(by_idx),
        "doc_stats": doc_stats(doc),
        "media_count": len(media_hashes(args.docx)),
    }
    (args.out / "stats.json").write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(stats, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
